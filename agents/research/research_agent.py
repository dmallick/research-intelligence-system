# src/agents/research/research_agent.py
import asyncio
import json
import os
from typing import Dict, Any, List, Optional, Union
from datetime import datetime, timezone
import logging
import aiohttp
import aiofiles
from pathlib import Path
import hashlib
import tempfile

# Document processing imports
try:
    from playwright.async_api import async_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False
    logging.warning("Playwright not available - web scraping will be limited")

try:
    from bs4 import BeautifulSoup
    import requests
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    logging.warning("BeautifulSoup not available")

try:
    import PyPDF2
    import docx
    PDF_DOCX_AVAILABLE = True
except ImportError:
    PDF_DOCX_AVAILABLE = False
    logging.warning("PDF/DOCX processing not available")

#import agents.research.research_agent
#from agents.research.research_agent import ResearchAgent
from agents.base.agent import BaseAgent
from core.message_queue import Message


class ResearchAgent(BaseAgent):
    """
    Specialized agent for research tasks including:
    - Web scraping
    - Academic paper retrieval
    - News API integration
    - Document parsing
    """
    
    def __init__(self, agent_id: str = "research_agent"):
        super().__init__(agent_id, agent_type="research")
        
        # Configuration
        self.config = {
            "max_content_length": 1000000,  # 1MB
            "request_timeout": 30,
            "max_concurrent_requests": 5,
            "cache_duration": 3600,  # 1 hour
            "supported_formats": ["html", "pdf", "docx", "txt", "json"],
            "news_api_key": os.getenv("NEWS_API_KEY"),
            "arxiv_base_url": "http://export.arxiv.org/api/query",
        }
        
        # Initialize components
        self.session: Optional[aiohttp.ClientSession] = None
        self.cache_dir = Path(tempfile.gettempdir()) / "research_agent_cache"
        self.cache_dir.mkdir(exist_ok=True)
        
        # Semaphore for concurrent requests
        self.request_semaphore = asyncio.Semaphore(self.config["max_concurrent_requests"])
        
        # Statistics
        self.stats = {
            "total_requests": 0,
            "successful_requests": 0,
            "failed_requests": 0,
            "documents_processed": 0,
            "cache_hits": 0,
            "started_at": datetime.now(timezone.utc)
        }
        
        # Register research-specific message handlers
        self._message_handlers.update({
            "research_task": self._handle_research_task,
            "web_scrape": self._handle_web_scrape,
            "arxiv_search": self._handle_arxiv_search,
            "news_search": self._handle_news_search,
            "document_parse": self._handle_document_parse,
            "url_extract": self._handle_url_extract,
        })
    
    async def initialize(self):
        """Initialize the research agent"""
        await super().initialize()
        
        # Initialize HTTP session
        self.session = aiohttp.ClientSession(
            timeout=aiohttp.ClientTimeout(total=self.config["request_timeout"]),
            headers={
                "User-Agent": "ResearchAgent/1.0 (Multi-Agent Research System)"
            }
        )
        
        self.logger.info("🔬 Research Agent initialized with capabilities:")
        self.logger.info(f"  - Web scraping: {'✅' if PLAYWRIGHT_AVAILABLE else '⚠️  Limited'}")
        self.logger.info(f"  - HTML parsing: {'✅' if BEAUTIFULSOUP_AVAILABLE else '❌'}")
        self.logger.info(f"  - PDF/DOCX: {'✅' if PDF_DOCX_AVAILABLE else '❌'}")
        self.logger.info(f"  - News API: {'✅' if self.config['news_api_key'] else '⚠️  No API key'}")
    
    async def shutdown(self):
        """Shutdown the research agent"""
        if self.session:
            await self.session.close()
        await super().shutdown()
    
    async def execute_task(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """Execute a research task"""
        task_type = task.get("type", "unknown")
        task_id = task.get("task_id", "unknown")
        
        self.stats["total_requests"] += 1
        
        try:
            if task_type == "web_scrape":
                result = await self._web_scrape(task.get("url"), task.get("options", {}))
            elif task_type == "arxiv_search":
                result = await self._arxiv_search(task.get("query"), task.get("max_results", 10))
            elif task_type == "news_search":
                result = await self._news_search(task.get("query"), task.get("options", {}))
            elif task_type == "document_parse":
                result = await self._parse_document(task.get("document_path"))
            elif task_type == "url_batch_extract":
                result = await self._batch_url_extract(task.get("urls", []))
            elif task_type == "research_summary":
                result = await self._create_research_summary(task.get("sources", []))
            else:
                raise ValueError(f"Unknown research task type: {task_type}")
            
            self.stats["successful_requests"] += 1
            self.stats["documents_processed"] += 1
            
            return {
                "task_id": task_id,
                "status": "completed",
                "result": result,
                "processed_at": datetime.now(timezone.utc).isoformat(),
                "agent_id": self.agent_id
            }
            
        except Exception as e:
            self.stats["failed_requests"] += 1
            self.logger.error(f"Research task failed: {e}")
            
            return {
                "task_id": task_id,
                "status": "failed",
                "error": str(e),
                "processed_at": datetime.now(timezone.utc).isoformat(),
                "agent_id": self.agent_id
            }
    
    async def get_status(self) -> Dict[str, Any]:
        """Get research agent status"""
        uptime = datetime.now(timezone.utc) - self.stats["started_at"]
        
        return {
            "capabilities": {
                "web_scraping": PLAYWRIGHT_AVAILABLE,
                "html_parsing": BEAUTIFULSOUP_AVAILABLE,
                "document_parsing": PDF_DOCX_AVAILABLE,
                "news_api": bool(self.config["news_api_key"]),
                "arxiv_search": True
            },
            "statistics": {
                **self.stats,
                "uptime_seconds": uptime.total_seconds(),
                "success_rate": (
                    self.stats["successful_requests"] / max(self.stats["total_requests"], 1)
                ) * 100
            },
            "configuration": {
                "max_concurrent_requests": self.config["max_concurrent_requests"],
                "request_timeout": self.config["request_timeout"],
                "supported_formats": self.config["supported_formats"]
            }
        }
    
    # Message Handlers
    async def _handle_research_task(self, message: Message):
        """Handle general research task messages"""
        try:
            task_data = message.payload
            result = await self.execute_task(task_data)
            
            await self.send_message(
                message.from_agent,
                "research_result",
                result,
                correlation_id=message.correlation_id
            )
            
        except Exception as e:
            await self.send_message(
                message.from_agent,
                "research_error",
                {
                    "task": message.payload,
                    "error": str(e),
                    "agent_id": self.agent_id
                },
                correlation_id=message.correlation_id
            )
    
    async def _handle_web_scrape(self, message: Message):
        """Handle web scraping requests"""
        url = message.payload.get("url")
        options = message.payload.get("options", {})
        
        if not url:
            await self.send_message(
                message.from_agent,
                "scrape_error",
                {"error": "No URL provided"},
                correlation_id=message.correlation_id
            )
            return
        
        try:
            result = await self._web_scrape(url, options)
            
            await self.send_message(
                message.from_agent,
                "scrape_result",
                {
                    "url": url,
                    "result": result,
                    "scraped_at": datetime.now(timezone.utc).isoformat()
                },
                correlation_id=message.correlation_id
            )
            
        except Exception as e:
            await self.send_message(
                message.from_agent,
                "scrape_error",
                {
                    "url": url,
                    "error": str(e)
                },
                correlation_id=message.correlation_id
            )
    
    async def _handle_arxiv_search(self, message: Message):
        """Handle arXiv search requests"""
        query = message.payload.get("query")
        max_results = message.payload.get("max_results", 10)
        
        if not query:
            await self.send_message(
                message.from_agent,
                "arxiv_error",
                {"error": "No search query provided"},
                correlation_id=message.correlation_id
            )
            return
        
        try:
            result = await self._arxiv_search(query, max_results)
            
            await self.send_message(
                message.from_agent,
                "arxiv_result",
                {
                    "query": query,
                    "papers": result,
                    "searched_at": datetime.now(timezone.utc).isoformat()
                },
                correlation_id=message.correlation_id
            )
            
        except Exception as e:
            await self.send_message(
                message.from_agent,
                "arxiv_error",
                {
                    "query": query,
                    "error": str(e)
                },
                correlation_id=message.correlation_id
            )
    
    async def _handle_news_search(self, message: Message):
        """Handle news search requests"""
        query = message.payload.get("query")
        options = message.payload.get("options", {})
        
        if not query:
            await self.send_message(
                message.from_agent,
                "news_error",
                {"error": "No search query provided"},
                correlation_id=message.correlation_id
            )
            return
        
        try:
            result = await self._news_search(query, options)
            
            await self.send_message(
                message.from_agent,
                "news_result",
                {
                    "query": query,
                    "articles": result,
                    "searched_at": datetime.now(timezone.utc).isoformat()
                },
                correlation_id=message.correlation_id
            )
            
        except Exception as e:
            await self.send_message(
                message.from_agent,
                "news_error",
                {
                    "query": query,
                    "error": str(e)
                },
                correlation_id=message.correlation_id
            )
    
    async def _handle_document_parse(self, message: Message):
        """Handle document parsing requests"""
        document_path = message.payload.get("document_path")
        
        if not document_path:
            await self.send_message(
                message.from_agent,
                "parse_error",
                {"error": "No document path provided"},
                correlation_id=message.correlation_id
            )
            return
        
        try:
            result = await self._parse_document(document_path)
            
            await self.send_message(
                message.from_agent,
                "parse_result",
                {
                    "document_path": document_path,
                    "content": result,
                    "parsed_at": datetime.now(timezone.utc).isoformat()
                },
                correlation_id=message.correlation_id
            )
            
        except Exception as e:
            await self.send_message(
                message.from_agent,
                "parse_error",
                {
                    "document_path": document_path,
                    "error": str(e)
                },
                correlation_id=message.correlation_id
            )
    
    async def _handle_url_extract(self, message: Message):
        """Handle URL content extraction requests"""
        urls = message.payload.get("urls", [])
        
        if not urls:
            await self.send_message(
                message.from_agent,
                "extract_error",
                {"error": "No URLs provided"},
                correlation_id=message.correlation_id
            )
            return
        
        try:
            result = await self._batch_url_extract(urls)
            
            await self.send_message(
                message.from_agent,
                "extract_result",
                {
                    "results": result,
                    "extracted_at": datetime.now(timezone.utc).isoformat()
                },
                correlation_id=message.correlation_id
            )
            
        except Exception as e:
            await self.send_message(
                message.from_agent,
                "extract_error",
                {
                    "urls": urls,
                    "error": str(e)
                },
                correlation_id=message.correlation_id
            )
    
    # Core Research Methods
    async def _web_scrape(self, url: str, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """Scrape web content from a URL"""
        options = options or {}
        
        # Check cache first
        cache_key = self._get_cache_key(url, options)
        cached_result = await self._get_from_cache(cache_key)
        if cached_result:
            self.stats["cache_hits"] += 1
            return cached_result
        
        async with self.request_semaphore:
            if PLAYWRIGHT_AVAILABLE and options.get("use_playwright", False):
                result = await self._scrape_with_playwright(url, options)
            else:
                result = await self._scrape_with_requests(url, options)
        
        # Cache the result
        await self._save_to_cache(cache_key, result)
        return result
    
    async def _scrape_with_playwright(self, url: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Scrape using Playwright for JavaScript-heavy sites"""
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            try:
                page = await browser.new_page()
                
                # Set viewport and user agent
                await page.set_viewport_size({"width": 1920, "height": 1080})
                
                # Navigate to page
                response = await page.goto(url, wait_until="networkidle")
                
                if not response or response.status >= 400:
                    raise Exception(f"Failed to load page: HTTP {response.status if response else 'unknown'}")
                
                # Wait for specific selector if provided
                if options.get("wait_for_selector"):
                    await page.wait_for_selector(options["wait_for_selector"], timeout=10000)
                
                # Get content
                content = await page.content()
                title = await page.title()
                
                # Extract text content
                text_content = await page.evaluate("""
                    () => {
                        // Remove script and style elements
                        const scripts = document.querySelectorAll('script, style');
                        scripts.forEach(el => el.remove());
                        
                        return document.body.innerText || document.body.textContent || '';
                    }
                """)
                
                return {
                    "url": url,
                    "title": title,
                    "html_content": content,
                    "text_content": text_content,
                    "status_code": response.status,
                    "method": "playwright",
                    "scraped_at": datetime.now(timezone.utc).isoformat()
                }
                
            finally:
                await browser.close()
    
    async def _scrape_with_requests(self, url: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Scrape using aiohttp and BeautifulSoup"""
        async with self.session.get(url) as response:
            if response.status >= 400:
                raise Exception(f"HTTP {response.status}: {response.reason}")
            
            content = await response.text()
            
            if BEAUTIFULSOUP_AVAILABLE:
                soup = BeautifulSoup(content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                title = soup.find('title')
                title_text = title.get_text().strip() if title else ""
                
                text_content = soup.get_text()
                # Clean up whitespace
                lines = (line.strip() for line in text_content.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text_content = ' '.join(chunk for chunk in chunks if chunk)
                
            else:
                title_text = ""
                text_content = content
            
            return {
                "url": url,
                "title": title_text,
                "html_content": content,
                "text_content": text_content,
                "status_code": response.status,
                "method": "aiohttp",
                "scraped_at": datetime.now(timezone.utc).isoformat()
            }
    
    async def _arxiv_search(self, query: str, max_results: int = 10) -> List[Dict[str, Any]]:
        """Search arXiv for academic papers"""
        params = {
            'search_query': query,
            'max_results': max_results,
            'sortBy': 'relevance',
            'sortOrder': 'descending'
        }
        
        async with self.session.get(self.config["arxiv_base_url"], params=params) as response:
            if response.status >= 400:
                raise Exception(f"arXiv API error: HTTP {response.status}")
            
            content = await response.text()
            
            papers = []
            if BEAUTIFULSOUP_AVAILABLE:
                soup = BeautifulSoup(content, 'xml')
                
                for entry in soup.find_all('entry'):
                    paper = {
                        'title': entry.find('title').get_text().strip() if entry.find('title') else "",
                        'authors': [
                            author.find('name').get_text().strip() 
                            for author in entry.find_all('author')
                            if author.find('name')
                        ],
                        'summary': entry.find('summary').get_text().strip() if entry.find('summary') else "",
                        'published': entry.find('published').get_text().strip() if entry.find('published') else "",
                        'updated': entry.find('updated').get_text().strip() if entry.find('updated') else "",
                        'arxiv_id': entry.find('id').get_text().split('/')[-1] if entry.find('id') else "",
                        'categories': [
                            cat.get('term') for cat in entry.find_all('category')
                            if cat.get('term')
                        ],
                        'pdf_url': None
                    }
                    
                    # Find PDF link
                    for link in entry.find_all('link'):
                        if link.get('type') == 'application/pdf':
                            paper['pdf_url'] = link.get('href')
                            break
                    
                    papers.append(paper)
            
            return papers
    
    async def _news_search(self, query: str, options: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Search for news articles using News API"""
        if not self.config["news_api_key"]:
            raise Exception("News API key not configured")
        
        options = options or {}
        
        params = {
            'q': query,
            'apiKey': self.config["news_api_key"],
            'pageSize': options.get('max_results', 20),
            'sortBy': options.get('sort_by', 'relevancy'),
            'language': options.get('language', 'en')
        }
        
        # Add date range if provided
        if options.get('from_date'):
            params['from'] = options['from_date']
        if options.get('to_date'):
            params['to'] = options['to_date']
        
        url = "https://newsapi.org/v2/everything"
        
        async with self.session.get(url, params=params) as response:
            if response.status >= 400:
                error_text = await response.text()
                raise Exception(f"News API error: HTTP {response.status} - {error_text}")
            
            data = await response.json()
            
            if data['status'] != 'ok':
                raise Exception(f"News API error: {data.get('message', 'Unknown error')}")
            
            articles = []
            for article in data.get('articles', []):
                articles.append({
                    'title': article.get('title', ''),
                    'description': article.get('description', ''),
                    'content': article.get('content', ''),
                    'url': article.get('url', ''),
                    'source': article.get('source', {}).get('name', ''),
                    'author': article.get('author', ''),
                    'published_at': article.get('publishedAt', ''),
                    'url_to_image': article.get('urlToImage', '')
                })
            
            return articles
    
    async def _parse_document(self, document_path: str) -> Dict[str, Any]:
        """Parse various document formats"""
        file_path = Path(document_path)
        
        if not file_path.exists():
            raise Exception(f"Document not found: {document_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf' and PDF_DOCX_AVAILABLE:
            return await self._parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc'] and PDF_DOCX_AVAILABLE:
            return await self._parse_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            return await self._parse_text(file_path)
        elif file_extension in ['.html', '.htm']:
            return await self._parse_html(file_path)
        elif file_extension == '.json':
            return await self._parse_json(file_path)
        else:
            raise Exception(f"Unsupported document format: {file_extension}")
    
    async def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF document"""
        text_content = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                    }
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        self.logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {e}")
        
        return {
            'file_path': str(file_path),
            'file_type': 'pdf',
            'content': text_content.strip(),
            'metadata': metadata,
            'page_count': len(pdf_reader.pages) if 'pdf_reader' in locals() else 0,
            'parsed_at': datetime.now(timezone.utc).isoformat()
        }
    
    async def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text content
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'last_modified_by': doc.core_properties.last_modified_by or ''
            }
            
            return {
                'file_path': str(file_path),
                'file_type': 'docx',
                'content': '\n'.join(text_content),
                'metadata': metadata,
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                'parsed_at': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {e}")
    
    async def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text document"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            return {
                'file_path': str(file_path),
                'file_type': 'text',
                'content': content,
                'metadata': {
                    'size_bytes': file_path.stat().st_size,
                    'lines': len(content.splitlines()),
                    'characters': len(content)
                },
                'parsed_at': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            raise Exception(f"Failed to parse text file: {e}")
    
    async def _parse_html(self, file_path: Path) -> Dict[str, Any]:
        """Parse HTML document"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                html_content = await file.read()
            
            text_content = html_content
            title = ""
            
            if BEAUTIFULSOUP_AVAILABLE:
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                title_tag = soup.find('title')
                title = title_tag.get_text().strip() if title_tag else ""
                
                text_content = soup.get_text()
                # Clean up whitespace
                lines = (line.strip() for line in text_content.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text_content = ' '.join(chunk for chunk in chunks if chunk)
            
            return {
                'file_path': str(file_path),
                'file_type': 'html',
                'content': text_content,
                'html_content': html_content,
                'metadata': {
                    'title': title,
                    'size_bytes': file_path.stat().st_size
                },
                'parsed_at': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            raise Exception(f"Failed to parse HTML file: {e}")
    
    async def _parse_json(self, file_path: Path) -> Dict[str, Any]:
        """Parse JSON document"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            json_data = json.loads(content)
            
            return {
                'file_path': str(file_path),
                'file_type': 'json',
                'content': json.dumps(json_data, indent=2),
                'json_data': json_data,
                'metadata': {
                    'size_bytes': file_path.stat().st_size,
                    'keys': list(json_data.keys()) if isinstance(json_data, dict) else None,
                    'items_count': len(json_data) if isinstance(json_data, (list, dict)) else None
                },
                'parsed_at': datetime.now(timezone.utc).isoformat()
            }
            
        except Exception as e:
            raise Exception(f"Failed to parse JSON file: {e}")
    
    async def _batch_url_extract(self, urls: List[str]) -> List[Dict[str, Any]]:
        """Extract content from multiple URLs concurrently"""
        semaphore = asyncio.Semaphore(self.config["max_concurrent_requests"])
        
        async def extract_single_url(url: str) -> Dict[str, Any]:
            async with semaphore:
                try:
                    result = await self._web_scrape(url)
                    return {
                        "url": url,
                        "status": "success",
                        "data": result
                    }
                except Exception as e:
                    return {
                        "url": url,
                        "status": "error",
                        "error": str(e)
                    }
        
        # Execute all extractions concurrently
        tasks = [extract_single_url(url) for url in urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Handle any exceptions from gather
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append({
                    "url": urls[i],
                    "status": "error",
                    "error": str(result)
                })
            else:
                processed_results.append(result)
        
        return processed_results
    
    async def _create_research_summary(self, sources: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Create a research summary from multiple sources"""
        summary = {
            "total_sources": len(sources),
            "successful_sources": 0,
            "failed_sources": 0,
            "content_summary": {},
            "key_findings": [],
            "sources_by_type": {},
            "generated_at": datetime.now(timezone.utc).isoformat()
        }
        
        all_text = []
        
        for source in sources:
            if source.get("status") == "success":
                summary["successful_sources"] += 1
                
                source_type = source.get("type", "unknown")
                if source_type not in summary["sources_by_type"]:
                    summary["sources_by_type"][source_type] = 0
                summary["sources_by_type"][source_type] += 1
                
                # Extract text content
                data = source.get("data", {})
                text = data.get("text_content") or data.get("content", "")
                if text:
                    all_text.append(text)
            else:
                summary["failed_sources"] += 1
        
        # Create basic content analysis
        if all_text:
            combined_text = " ".join(all_text)
            summary["content_summary"] = {
                "total_characters": len(combined_text),
                "total_words": len(combined_text.split()),
                "average_words_per_source": len(combined_text.split()) / len(all_text) if all_text else 0
            }
            
            # Simple keyword extraction (most frequent words)
            words = combined_text.lower().split()
            word_freq = {}
            for word in words:
                if len(word) > 3:  # Skip short words
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # Get top 10 most frequent words
            top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
            summary["key_findings"] = [{"word": word, "frequency": freq} for word, freq in top_words]
        
        return summary
    
    # Utility methods
    def _get_cache_key(self, url: str, options: Dict[str, Any]) -> str:
        """Generate cache key for URL and options"""
        cache_string = f"{url}:{json.dumps(options, sort_keys=True)}"
        return hashlib.md5(cache_string.encode()).hexdigest()
    
    async def _get_from_cache(self, cache_key: str) -> Optional[Dict[str, Any]]:
        """Get result from cache if available and not expired"""
        cache_file = self.cache_dir / f"{cache_key}.json"
        
        if not cache_file.exists():
            return None
        
        try:
            # Check if cache is expired
            file_age = datetime.now().timestamp() - cache_file.stat().st_mtime
            if file_age > self.config["cache_duration"]:
                cache_file.unlink()  # Remove expired cache
                return None
            
            async with aiofiles.open(cache_file, 'r') as f:
                content = await f.read()
                return json.loads(content)
                
        except Exception as e:
            self.logger.warning(f"Failed to read cache {cache_key}: {e}")
            return None
    
    async def _save_to_cache(self, cache_key: str, data: Dict[str, Any]):
        """Save result to cache"""
        cache_file = self.cache_dir / f"{cache_key}.json"
        
        try:
            async with aiofiles.open(cache_file, 'w') as f:
                await f.write(json.dumps(data, indent=2, default=str))
        except Exception as e:
            self.logger.warning(f"Failed to save cache {cache_key}: {e}")
    
    async def clear_cache(self):
        """Clear all cached results"""
        try:
            for cache_file in self.cache_dir.glob("*.json"):
                cache_file.unlink()
            self.logger.info("🧹 Research agent cache cleared")
        except Exception as e:
            self.logger.error(f"Failed to clear cache: {e}")
    
    async def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics"""
        try:
            cache_files = list(self.cache_dir.glob("*.json"))
            total_size = sum(f.stat().st_size for f in cache_files)
            
            return {
                "cache_files": len(cache_files),
                "total_size_bytes": total_size,
                "total_size_mb": round(total_size / (1024 * 1024), 2),
                "cache_directory": str(self.cache_dir),
                "cache_hits": self.stats["cache_hits"]
            }
        except Exception as e:
            return {"error": str(e)}


# Research Agent Factory
def create_research_agent(agent_id: str = None) -> ResearchAgent:
    """Create a new research agent instance"""
    if agent_id is None:
        agent_id = f"research_agent_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
    return ResearchAgent(agent_id)


# Research Task Templates
class ResearchTaskTemplates:
    """Pre-defined research task templates"""
    
    @staticmethod
    def web_scrape_task(url: str, use_playwright: bool = False) -> Dict[str, Any]:
        """Template for web scraping task"""
        return {
            "type": "web_scrape",
            "url": url,
            "options": {
                "use_playwright": use_playwright,
                "wait_for_selector": None
            }
        }
    
    @staticmethod
    def arxiv_search_task(query: str, max_results: int = 10) -> Dict[str, Any]:
        """Template for arXiv search task"""
        return {
            "type": "arxiv_search",
            "query": query,
            "max_results": max_results
        }
    
    @staticmethod
    def news_search_task(
        query: str, 
        max_results: int = 20, 
        language: str = "en",
        sort_by: str = "relevancy"
    ) -> Dict[str, Any]:
        """Template for news search task"""
        return {
            "type": "news_search",
            "query": query,
            "options": {
                "max_results": max_results,
                "language": language,
                "sort_by": sort_by
            }
        }
    
    @staticmethod
    def document_parse_task(document_path: str) -> Dict[str, Any]:
        """Template for document parsing task"""
        return {
            "type": "document_parse",
            "document_path": document_path
        }
    
    @staticmethod
    def batch_url_extract_task(urls: List[str]) -> Dict[str, Any]:
        """Template for batch URL extraction task"""
        return {
            "type": "url_batch_extract",
            "urls": urls
        }
    
    @staticmethod
    def research_summary_task(sources: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Template for research summary task"""
        return {
            "type": "research_summary",
            "sources": sources
        }